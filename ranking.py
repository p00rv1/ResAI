import os
import re
import numpy as np
import pandas as pd
from typing import List, Dict, Any
from docx import Document

def rank_candidates(
    input_file="top5000_candidates.csv",
    jd_file="job_description.docx",
    output_file="top100.csv"
):

    df = pd.read_csv(input_file)

    doc = Document(jd_file)

    jd_text = "\n".join(
        para.text
        for para in doc.paragraphs
    )

    jd_lower = jd_text.lower()
    
    all_skills = set()

    for s in df["skills_text"]:
        all_skills.update(s)
    doc = Document("job_description.docx")

    jd_text = "\n".join(
        para.text
        for para in doc.paragraphs
    )
    jd_lower = jd_text.lower()

    jd_skills = set()

    for skill in all_skills:
        if skill in jd_lower:
            jd_skills.add(skill)
    def skill_overlap(candidate_skills, jd_skills):
        sk = set()
        sk.update(candidate_skills.lower().split())
        if len(jd_skills) == 0:
            return 0

        matched = sk & jd_skills

        return len(matched) / len(jd_skills)
    df["skill_overlap"] = df["skills_text"].apply(
        lambda x: skill_overlap(x, jd_skills)
    )
    # --------------------------------------------------
    # STEP 1: Find all assessment columns automatically
    # --------------------------------------------------

    assessment_cols = [
        col
        for col in df.columns
        if "skill_assessment_scores" in col.lower()
    ]


    # --------------------------------------------------
    # STEP 2: Find which assessment topics appear in JD
    # --------------------------------------------------

    jd_lower = jd_text.lower()

    relevant_assessment_cols = []

    for col in assessment_cols:

        # Example:
        # redrob_signals_skill_assessment_scores_MLOps
        # -> mlops

        topic = col.split("scores_")[-1].lower()

        if topic in jd_lower:
    
            relevant_assessment_cols.append(col)



    # --------------------------------------------------
    # STEP 3: Create assessment feature
    # --------------------------------------------------

    def compute_assessment_feature(row):

        scores = []

        # Count how many JD-relevant assessments
        # the candidate actually has

        for col in relevant_assessment_cols:

            value = row[col]

            if pd.notna(value):
                scores.append(value)

        # No matching assessments

        if len(scores) == 0:
            return 0

        # Average score quality
        mean_score = np.mean(scores)

        # Coverage of JD skills
        coverage = (
            len(scores)
            / len(relevant_assessment_cols)
        )

        # Final assessment feature
        # Quality + Coverage

        feature = (
            0.7 * mean_score +
            0.3 * coverage * 100
        )

        return feature


    # --------------------------------------------------
    # STEP 4: Apply to candidates
    # --------------------------------------------------

    df["assessment_feature"] = (
        df.apply(
            compute_assessment_feature,
            axis=1
        )
    )

    print(
        df[
            ["assessment_feature"]
        ]
        .sort_values(
            "assessment_feature",
            ascending=False
        )
        .head()
    )

    jd_words = set(
        jd_text.lower().split()
    )
    def title_similarity(title):

        if pd.isna(title):
            return 0

        title_words = set(
            str(title).lower().split()
        )

        overlap = (
            len(
                title_words & jd_words
            )
        )

        return overlap

    df["title_similarity"] = (
        df["profile_current_title"]
        .apply(title_similarity)
    )


    redrob_cols = [
        "redrob_signals_saved_by_recruiters_30d",
        "redrob_signals_search_appearance_30d",
        "redrob_signals_profile_views_received_30d",
        "redrob_signals_recruiter_response_rate",
        "redrob_signals_interview_completion_rate",
        "redrob_signals_offer_acceptance_rate",
        "redrob_signals_profile_completeness_score",
        "redrob_signals_github_activity_score",
        "redrob_signals_connection_count",
        "redrob_signals_applications_submitted_30d"
    ]

    for col in redrob_cols:
        df[col] = df[col].fillna(0)

    from sklearn.preprocessing import MinMaxScaler

    scaler = MinMaxScaler()

    df[redrob_cols] = scaler.fit_transform(
        df[redrob_cols]
    )

    bool_cols = [
        "redrob_signals_verified_email",
        "redrob_signals_verified_phone",
        "redrob_signals_linkedin_connected",
        "redrob_signals_open_to_work_flag",
        "redrob_signals_willing_to_relocate"
    ]

    for col in bool_cols:
        df[col] = (
            df[col]
            .fillna(False)
            .astype(int)
        )

    df["redrob_quality_score"] = (

        # recruiter demand
        0.20 * df["redrob_signals_saved_by_recruiters_30d"] +
        0.15 * df["redrob_signals_search_appearance_30d"] +
        0.10 * df["redrob_signals_profile_views_received_30d"] +
        0.10 * df["redrob_signals_recruiter_response_rate"] +

        # candidate quality
        0.10 * df["redrob_signals_interview_completion_rate"] +
        0.10 * df["redrob_signals_offer_acceptance_rate"] +
        0.05 * df["redrob_signals_profile_completeness_score"] +

        # activity
        0.05 * df["redrob_signals_github_activity_score"] +
        0.05 * df["redrob_signals_connection_count"] +
        0.03 * df["redrob_signals_applications_submitted_30d"] +

        # trust
        0.02 * df["redrob_signals_verified_email"] +
        0.02 * df["redrob_signals_verified_phone"] +
        0.02 * df["redrob_signals_linkedin_connected"] +

        # availability
        0.01 * df["redrob_signals_open_to_work_flag"]
    )
    from datetime import datetime

    today = pd.Timestamp.today()

    df["last_active_days"] = (
        today - pd.to_datetime(
            df["redrob_signals_last_active_date"],
            errors="coerce"
        )
    ).dt.days
    df["last_active_score"] = (
        1 / (1 + df["last_active_days"].fillna(999))
    )

    df["ye"]=df["profile_years_of_experience"].fillna(0)
    from sklearn.preprocessing import MinMaxScaler

    features = [
        "bm25_score",
        "skill_strength",
        "skill_overlap",
        "assessment_feature",
        "title_similarity",
        "redrob_quality_score",
        "last_active_score",
        "ye"
    ]

    scaler = MinMaxScaler()

    df[features] = scaler.fit_transform(
        df[features]
    )
    df["final_score"] = (

        0.22 * df["bm25_score"]

        + 0.16 * df["skill_strength"]

        + 0.16 * df["skill_overlap"]

        + 0.18 * df["assessment_feature"]

        + 0.04 * df["title_similarity"]

        + 0.20 * df["redrob_quality_score"]

        + 0.03 * df["last_active_score"]
        +
        0.01 * df["ye"]

    )
    # Sort by final score
    final_ranked = (
        df
        .sort_values(
            "final_score",
            ascending=False
        )
        .reset_index(drop=True)
    )

    # Assign ranks after sorting
    final_ranked["rank"] = final_ranked.index + 1

    # Take top 100
    top100 = final_ranked.head(100)

    top100.to_csv(
    output_file,
    index=False
    )

    print(
        f"Saved {output_file}"
    )

    return top100
if __name__ == "__main__":
    rank_candidates()