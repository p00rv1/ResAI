import os
import re
import numpy as np
import pandas as pd
from typing import List, Dict, Any
from docx import Document
from rank_bm25 import BM25Okapi
from sentence_transformers import SentenceTransformer

class FastFilteredSignalsRanker:
    def __init__(self, model_name: str = "all-MiniLM-L6-v2"):
        """Engine optimized to execute hard data constraints prior to text embedding generation."""
        print("Initializing local semantic encoder...")
        self.encoder = SentenceTransformer(model_name)
        self.df: pd.DataFrame = None

    @staticmethod
    def read_docx_job_description(file_path: str) -> str:
        """Extracts text metrics from a .docx file safely."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Job Description file at '{file_path}' was not found.")
        doc = Document(file_path)
        full_text = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n".join(full_text)

    def load_parquet_data(self, file_path: str):
        """Loads parquet data frames straight to RAM without performing expensive processing steps."""
        print(f"Loading raw table from {file_path}...")
        self.df = pd.read_parquet(file_path)
        print(f"Successfully parked {len(self.df):,} rows in system memory.")

    def _parse_jd_boundaries(self, jd_text: str) -> Dict[str, Any]:
        """Parses experience constraints and salary expectations from raw text."""
        exp_matches = re.findall(r'(\d+)\s*[\-\+]?\s*(?:to\s*(\d+))?\s*(?:years|yrs)', jd_text, re.IGNORECASE)
        sal_matches = re.findall(r'(\d+)\s*[\-\+]?\s*(?:to\s*(\d+))?\s*(?:lpa|lakhs)', jd_text, re.IGNORECASE)
        
        return {
            "target_exp": float(exp_matches[0][0]) if exp_matches else 4.0,
            "max_salary_cap": float(sal_matches[0][1]) if (sal_matches and len(sal_matches[0]) > 1 and sal_matches[0][1]) else 999.0
        }

    def rank(self, job_description_text: str, top_k: int = 100) -> List[Dict[str, Any]]:
        """Executes a pipeline that runs hard filtering first, saving tokenization and embeddings for last."""
        constraints = self._parse_jd_boundaries(job_description_text)
        tokenized_jd = job_description_text.lower().split()
        
        # ==============================================================================
        # STEP 1: REMOVE HONEYPOTS & HARD UNVIABLE KNOCKOUTS (Vectorized Pandas)
        # ==============================================================================
        print("Executing Phase 1 hard constraint check across complete database...")
        
        # Build logical boolean arrays for high-velocity masking
        is_not_honeypot = (
            (self.df['profile_completeness_score'].fillna(0) >= 50) & 
            (self.df['recruiter_response_rate'].fillna(0.5) > 0.1) &
            (self.df['search_appearance_30d'].fillna(0) >= 0)
        )
        
        is_financially_viable = (
            self.df['expected_salary_min_lpa'].fillna(0) <= (constraints["max_salary_cap"] * 1.20)
        )
        
        # Slice original data safely using our boolean indicators
        filtered_df = self.df[is_not_honeypot & is_financially_viable].copy()
        current_pool_size = len(filtered_df)
        print(f"Filtered pool size down from {len(self.df):,} to {current_pool_size:,} candidates.")
        
        if current_pool_size == 0:
            print("Warning: Hard criteria knocked out all entries. Softening fallback to base pool.")
            filtered_df = self.df.copy()
            current_pool_size = len(filtered_df)

        # ==============================================================================
        # STEP 2: LOCAL HIGH-PASS LEXICAL SELECTION (BM25)
        # ==============================================================================
        print("Building local temporary BM25 maps over clean subset candidates...")
        text_corpus_series = filtered_df['master_text_profile'].fillna("").astype(str).str.lower()
        tokenized_corpus = [doc.split() for doc in text_corpus_series.values]
        
        local_bm25 = BM25Okapi(tokenized_corpus)
        bm25_scores = local_bm25.get_scores(tokenized_jd)
        
        # Isolate the top few thousand for deep semantic verification
        stage1_budget = min(3000, current_pool_size)
        top_stage1_local_indices = np.argsort(bm25_scores)[::-1][:stage1_budget]
        
        # Isolate the target subset data slice
        semantic_target_df = filtered_df.iloc[top_stage1_local_indices].copy()

        # ==============================================================================
        # STEP 3: DENSE EMBEDDING GENERATION ON FILTERED SUBSET ONLY (Fast Runtime)
        # ==============================================================================
        print(f"Generating semantic vector pairs for only {len(semantic_target_df)} candidates...")
        target_text_list = semantic_target_df['master_text_profile'].fillna("").astype(str).str.lower().tolist()
        
        # Compute subset matrix instantly
        subset_embeddings = self.encoder.encode(
            target_text_list,
            batch_size=128,
            show_progress_bar=False,
            convert_to_numpy=True
        )
        
        jd_vector = self.encoder.encode(job_description_text, convert_to_numpy=True)
        
        # Run subset matrix cosine calculations
        jd_norm = np.linalg.norm(jd_vector)
        subset_norms = np.linalg.norm(subset_embeddings, axis=1)
        cosine_sims = np.dot(subset_embeddings, jd_vector) / (subset_norms * jd_norm + 1e-9)
        
        # Calculate base text compatibility metric
        sub_bm25_scores = bm25_scores[top_stage1_local_indices]
        norm_bm25 = (sub_bm25_scores - sub_bm25_scores.min()) / (sub_bm25_scores.max() - sub_bm25_scores.min() + 1e-9)
        text_scores = (0.4 * norm_bm25) + (0.6 * cosine_sims)
        
        # Narrow the pool down to the top 500
        stage2_budget = min(500, len(semantic_target_df))
        top_stage2_relative_indices = np.argsort(text_scores)[::-1][:stage2_budget]
        
        final_processing_df = semantic_target_df.iloc[top_stage2_relative_indices].copy()
        final_processing_df['score_text_base'] = text_scores[top_stage2_relative_indices]

        # ==============================================================================
        # STEP 4: VECTORIZED MULTI-VARIABLE FINAL SCORE
        # ==============================================================================
        print("Executing absolute weight balancing across candidate records...")
        
        # Guard missing metrics
        final_processing_df['years_of_experience'] = final_processing_df['years_of_experience'].fillna(0).astype(float)
        final_processing_df['avg_verified_skill_score'] = final_processing_df['avg_verified_skill_score'].fillna(0).astype(float)
        final_processing_df['github_activity_score'] = final_processing_df['github_activity_score'].fillna(-1).astype(float)
        final_processing_df['interview_completion_rate'] = final_processing_df['interview_completion_rate'].fillna(0).astype(float)
        final_processing_df['notice_period_days'] = final_processing_df['notice_period_days'].fillna(90).astype(float)
        
        # Experience Check
        exp_diff = constraints["target_exp"] - final_processing_df['years_of_experience']
        final_processing_df['score_experience'] = np.where(exp_diff <= 0, 1.0, np.maximum(0.0, 1.0 - exp_diff * 0.2))
        final_processing_df['score_experience'] *= (1.0 - (final_processing_df['num_past_companies'].fillna(1).astype(float) / 20.0).clip(0, 0.3))
        
        # Technical Quality Check
        git_norm = np.where(final_processing_df['github_activity_score'] >= 0, final_processing_df['github_activity_score'] / 100.0, 0.4)
        test_norm = final_processing_df['avg_verified_skill_score'] / 100.0
        completion_norm = final_processing_df['interview_completion_rate']
        final_processing_df['score_validation'] = (0.4 * test_norm) + (0.3 * git_norm) + (0.3 * completion_norm)
        
        # Market Intent Check
        notice_norm = 1.0 - (final_processing_df['notice_period_days'] / 180.0).clip(0, 0.6)
        otw_norm = np.where(final_processing_df['open_to_work_flag'] == True, 1.0, 0.4)
        response_norm = final_processing_df['recruiter_response_rate'].fillna(0.5)
        final_processing_df['score_liquidity'] = (0.4 * otw_norm) + (0.3 * notice_norm) + (0.3 * response_norm)
        
        # Compute Master Equation Evaluation Points
        final_processing_df['final_composite_score'] = (0.50 * final_processing_df['score_text_base']) + \
                                                       (0.20 * final_processing_df['score_validation']) + \
                                                       (0.15 * final_processing_df['score_experience']) + \
                                                       (0.15 * final_processing_df['score_liquidity'])
        
        # Isolate Final Ranked Data Output Frame
        ranked_df = final_processing_df.sort_values(by='final_composite_score', ascending=False).head(top_k)

        # ==============================================================================
        # STEP 5: EVIDENCE-BASED DETERMINISTIC EXPLANATIONS
        # ==============================================================================
        results = []
        for _, row in ranked_df.iterrows():
            features = {
                "text": float(row['score_text_base']),
                "validation": float(row['score_validation']),
                "experience": float(row['score_experience']),
                "liquidity": float(row['score_liquidity'])
            }
            
            results.append({
                "candidate_id": row['candidate_id'],
                "anonymized_name": row['anonymized_name'],
                "composite_score": float(row['final_composite_score']),
                "explanation": self._generate_explanation(row, features),
                "metrics_dashboard": {
                    "current_title": row['current_title'],
                    "years_of_experience": float(row['years_of_experience']),
                    "expected_salary_min_lpa": float(row['expected_salary_min_lpa']),
                    "notice_period_days": int(row['notice_period_days']),
                    "avg_verified_skill_score": float(row['avg_verified_skill_score'])
                }
            })
            
        return results

    def _generate_explanation(self, row: pd.Series, features: Dict[str, float]) -> str:
        """Translates final evaluation weights into deterministic feedback strings."""
        proof_points = []
        if features["text"] > 0.65:
            proof_points.append(f"strong alignment with requirements as a {row['current_title']}")
        if row['avg_verified_skill_score'] > 75:
            proof_points.append(f"distinguished Redrob test performance (Avg: {int(row['avg_verified_skill_score'])})")
        if features["experience"] >= 0.85:
            proof_points.append(f"solid longevity track record across {int(row['years_of_experience'])} years of industry tenure")
        if row['open_to_work_flag'] and row['notice_period_days'] <= 30:
            proof_points.append(f"immediate market availability with a clean {int(row['notice_period_days'])}-day notice window")
            
        if not proof_points:
            proof_points.append("balanced behavioral scores across overall platform footprint vectors")

        explanation = "Selected due to " + ", ".join(proof_points[:-1]) + ", and " + proof_points[-1] + "."
        return explanation[0].upper() + explanation[1:]

# --- QUICK CALL SANITY WRAPPER ---
if __name__ == "__main__":
    engine = FastFilteredSignalsRanker()
    engine.load_parquet_data("clean_candidates.parquet")
    
    import time
    start = time.time()
    
    raw_jd = engine.read_docx_job_description("job_description.docx")
    top_100_shortlist = engine.rank(raw_jd, top_k=100)
    # Assuming 'top_100_shortlist' is the array returned from your engine.rank() call
# top_100_shortlist = engine.rank(raw_jd, top_k=100)

# 1. Flatten the inner 'metrics_dashboard' fields into clean columns
    flattened_results = []
    for rank_idx, candidate in enumerate(top_100_shortlist):
        row = {
            "Rank": rank_idx + 1,
            "Candidate ID": candidate["candidate_id"],
            "Anonymized Name": candidate["anonymized_name"],
            "Match Score": round(candidate["composite_score"], 4),
            "Reasoning": candidate["explanation"],
            # Pull everything from the metrics dashboard out into individual columns
            "Current Title": candidate["metrics_dashboard"]["current_title"],
            "Years of Experience": candidate["metrics_dashboard"]["years_of_experience"],
            "Expected Salary (LPA)": candidate["metrics_dashboard"]["expected_salary_min_lpa"],
            "Notice Period (Days)": candidate["metrics_dashboard"]["notice_period_days"],
            "Avg Assessment Score": candidate["metrics_dashboard"]["avg_verified_skill_score"]
        }
        flattened_results.append(row)

    # 2. Convert the list of rows into a clean Pandas DataFrame
    export_df = pd.DataFrame(flattened_results)

    # 3. Write out to a CSV file
    output_filename = "ranked_top_candidates.csv"
    export_df.to_csv(output_filename, index=False, encoding="utf-8")

    print(f"Successfully exported the top candidates to: {output_filename}")
    print(f"\nPipeline execution finished flawlessly in: {time.time() - start:.2f} seconds!")
    