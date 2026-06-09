import docx2txt
import re
import os
import re
import numpy as np
import pandas as pd
from typing import List, Dict, Any
from docx import Document
from rank_bm25 import BM25Okapi
from sentence_transformers import SentenceTransformer


def read_docx_job_description(file_path: str) -> str:
        """Extracts text metrics from a .docx file safely."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Job Description file at '{file_path}' was not found.")
        doc = Document(file_path)
        full_text = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n".join(full_text)

# Test our parser module instantly
if __name__ == "__main__":
    requirements = read_docx_job_description("job_description.docx")
    
    print("Extracted Job Description:" + requirements)
   